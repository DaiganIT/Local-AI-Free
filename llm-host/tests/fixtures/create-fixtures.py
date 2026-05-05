#!/usr/bin/env python3
"""Create test PDF fixtures for integration tests."""

from fpdf import FPDF
import os

FIXTURES_DIR = os.path.dirname(os.path.abspath(__file__))


def create_text_and_image_pdf():
    """Create a PDF with text content and an embedded image."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Test Report", ln=True)
    pdf.cell(0, 10, "This is a test PDF document.", ln=True)
    pdf.cell(0, 10, "It contains text and an image.", ln=True)

    # Add a small red square image using built-in image support
    # Create a small PNG image inline
    from PIL import Image as PILImage
    import io

    # Create a 20x20 red square PNG
    img = PILImage.new("RGB", (20, 20), color="red")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)

    pdf.image(buf, x=10, y=50, w=40, h=40)
    pdf.output(os.path.join(FIXTURES_DIR, "text-and-image.pdf"))


def create_text_only_pdf():
    """Create a PDF with only text (no images)."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Meeting Notes", ln=True)
    pdf.cell(0, 10, "This is a text-only PDF document.", ln=True)
    pdf.cell(0, 10, "No images here.", ln=True)
    pdf.output(os.path.join(FIXTURES_DIR, "text-only.pdf"))


def create_docx_file():
    """Create a DOCX file for testing non-PDF document extraction."""
    # markitdown can extract from DOCX - we'll create one if python-docx is available
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test DOCX document.")
        doc.add_paragraph("It has some text content.")
        doc.save(os.path.join(FIXTURES_DIR, "test-document.docx"))
    except ImportError:
        # Create a minimal DOCX manually (it's a ZIP file)
        import zipfile
        import io

        docx_path = os.path.join(FIXTURES_DIR, "test-document.docx")

        content_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>Test Document</w:t></w:r></w:p>
    <w:p><w:r><w:t>This is a test DOCX document.</w:t></w:r></w:p>
  </w:body>
</w:document>"""

        rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

        word_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>"""

        styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
</w:styles>"""

        content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats.wordprocessingml.document.main+xml"/>
</Types>"""

        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('[Content_Types].xml', content_types_xml)
            zf.writestr('_rels/.rels', rels_xml)
            zf.writestr('word/document.xml', content_xml)
            zf.writestr('word/_rels/document.xml.rels', word_rels_xml)
            zf.writestr('word/styles.xml', styles_xml)


if __name__ == "__main__":
    create_text_and_image_pdf()
    print(f"Created text-and-image.pdf")
    create_text_only_pdf()
    print(f"Created text-only.pdf")
    create_docx_file()
    print(f"Created test-document.docx")